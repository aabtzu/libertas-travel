"""File format parsers: ICS, JSON trip data, Excel, Word documents."""

from __future__ import annotations

import base64
import json
import re
from datetime import datetime
from typing import Any

from agents.common.categories import normalize_category

# Single source of truth for supported upload extensions, used by all upload
# handlers and the frontend (via LibertasUpload in main.js, which must stay in sync).
SUPPORTED_EXTENSIONS = [
    ".pdf",
    ".txt",
    ".csv",
    ".png",
    ".jpg",
    ".jpeg",
    ".html",
    ".htm",
    ".eml",
    ".ics",
    ".json",
    ".xlsx",
    ".xls",
    ".docx",
]

# Regex patterns for extracting coordinates from Google Maps URLs
_GOOGLE_MAPS_COORD_RE = re.compile(r"@(-?\d+\.\d+),(-?\d+\.\d+)")
_GOOGLE_MAPS_QUERY_RE = re.compile(r"[?&]q=(-?\d+\.\d+),(-?\d+\.\d+)")


def extract_file_content(file_data: bytes, ext: str) -> dict[str, Any]:
    """Extract content from an uploaded file for LLM processing.

    Returns a dict with one or more of:
      - "text": str         , text content to pass as a message
      - "image_data": str   , base64-encoded image (for vision models)
      - "media_type": str   , MIME type of image_data
      - "items": list       , pre-parsed items (ICS/JSON fast path)
      - "error": str        , error message if extraction failed
    """
    ext = ext.lower().lstrip(".")

    if ext in ("txt", "html", "htm", "eml", "csv", "kml"):
        try:
            return {"text": file_data.decode("utf-8")}
        except UnicodeDecodeError:
            return {"text": file_data.decode("latin-1")}

    if ext in ("png", "jpg", "jpeg", "gif", "webp"):
        img = base64.standard_b64encode(file_data).decode("utf-8")
        mime = f"image/{ext}" if ext != "jpg" else "image/jpeg"
        return {"image_data": img, "media_type": mime}

    if ext == "ics":
        try:
            items = _parse_ics_file(file_data)
            if items:
                return {"items": items}
        except Exception:
            pass
        try:
            return {"text": file_data.decode("utf-8")}
        except UnicodeDecodeError:
            return {"text": file_data.decode("latin-1")}

    if ext == "json":
        try:
            items = _parse_json_trip(file_data)
            if items:
                return {"items": items}
        except Exception:
            pass
        try:
            return {"text": file_data.decode("utf-8")}
        except UnicodeDecodeError:
            return {"error": "Invalid JSON file encoding"}

    if ext in ("xlsx", "xls"):
        try:
            return {"text": _parse_excel_to_text(file_data, ext)}
        except ImportError:
            return {"error": "Excel processing unavailable, install openpyxl"}
        except Exception as e:
            return {"error": f"Error reading Excel file: {e}"}

    if ext in ("docx", "doc"):
        try:
            return {"text": _parse_word_to_text(file_data, ext)}
        except ImportError:
            return {"error": "Word processing unavailable, install python-docx"}
        except Exception as e:
            return {"error": f"Error reading Word document: {e}"}

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=file_data, filetype="pdf")
            text = "\n".join(page.get_text() for page in doc)
            if text.strip():
                doc.close()
                return {"text": text}
            # Scanned PDF, fall back to image of first page
            pix = doc[0].get_pixmap(matrix=fitz.Matrix(2, 2))
            img = base64.standard_b64encode(pix.tobytes("png")).decode("utf-8")
            doc.close()
            return {"image_data": img, "media_type": "image/png"}
        except ImportError:
            return {"error": "PDF processing unavailable, install pymupdf"}
        except Exception as e:
            return {"error": f"Error reading PDF: {e}"}

    return {"error": f"Unsupported file type: .{ext}"}


def extract_coords_from_url(url: str | None) -> tuple[float, float] | None:
    """Extract lat/lng from a Google Maps URL, or return None."""
    if not url or not isinstance(url, str) or "google" not in url.lower():
        return None
    for pattern in (_GOOGLE_MAPS_COORD_RE, _GOOGLE_MAPS_QUERY_RE):
        m = pattern.search(url)
        if m:
            lat, lng = float(m.group(1)), float(m.group(2))
            if -90 <= lat <= 90 and -180 <= lng <= 180:
                return (lat, lng)
    return None


def _normalize_item(item: dict[str, Any]) -> dict[str, Any]:
    """Normalize an item dict to our expected format."""
    normalized = {}

    # Title can be in various fields
    normalized["title"] = (
        item.get("title")
        or item.get("name")
        or item.get("summary")
        or item.get("event")
        or "Untitled"
    )

    # Category normalization
    # normalize_category is the single source of truth, defined in agents/common/categories.py
    cat = item.get("category") or item.get("type") or ""
    normalized["category"] = normalize_category(cat)

    # Date handling
    date = item.get("date") or item.get("start_date") or item.get("startDate")
    if date:
        if isinstance(date, str):
            if len(date) >= 10 and date[4] == "-":
                normalized["date"] = date[:10]
            else:
                for fmt in ["%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y"]:
                    try:
                        parsed = datetime.strptime(date[:10], fmt)
                        normalized["date"] = parsed.strftime("%Y-%m-%d")
                        break
                    except ValueError:
                        pass

    # Time handling
    time = item.get("time") or item.get("start_time") or item.get("startTime")
    if time:
        if isinstance(time, str):
            if ":" in time and len(time) >= 5:
                normalized["time"] = time[:5]
            elif len(time) == 4 and time.isdigit():
                normalized["time"] = f"{time[:2]}:{time[2:]}"

    # End time handling (for flights, trains, etc.)
    end_time = item.get("end_time") or item.get("endTime") or item.get("arrival_time")
    if end_time:
        if isinstance(end_time, str):
            if ":" in end_time and len(end_time) >= 5:
                normalized["end_time"] = end_time[:5]
            elif len(end_time) == 4 and end_time.isdigit():
                normalized["end_time"] = f"{end_time[:2]}:{end_time[2:]}"

    # Location
    loc = item.get("location")
    if isinstance(loc, dict):
        normalized["location"] = loc.get("name") or loc.get("address") or loc.get("city")
    elif isinstance(loc, str):
        normalized["location"] = loc
    elif item.get("city"):
        normalized["location"] = item.get("city")
    elif item.get("address"):
        normalized["location"] = item.get("address")

    # Notes
    notes = item.get("notes") or item.get("description") or item.get("details")
    if notes:
        normalized["notes"] = str(notes)[:500]

    # Day number if present
    if item.get("day") or item.get("day_number"):
        normalized["day"] = item.get("day") or item.get("day_number")

    # Coordinates, Google Maps exports use lat/latitude, lng/longitude/lon
    lat = item.get("latitude") or item.get("lat")
    lng = item.get("longitude") or item.get("lng") or item.get("lon")

    # Also check if location was a dict with coordinate sub-fields
    if (lat is None or lng is None) and isinstance(loc, dict):
        lat = lat or loc.get("latitude") or loc.get("lat")
        lng = lng or loc.get("longitude") or loc.get("lng") or loc.get("lon")

    # Validate both present, numeric, and in range
    if lat is not None and lng is not None:
        try:
            lat_f, lng_f = float(lat), float(lng)
            if -90 <= lat_f <= 90 and -180 <= lng_f <= 180:
                normalized["latitude"] = lat_f
                normalized["longitude"] = lng_f
        except (ValueError, TypeError):
            pass

    # Fallback: extract coordinates from Google Maps URLs in url/link fields
    if "latitude" not in normalized:
        for url_field in ("url", "website", "link", "google_maps_url"):
            coords = extract_coords_from_url(item.get(url_field))
            if coords:
                normalized["latitude"], normalized["longitude"] = coords
                break

    return normalized


def _guess_event_category(title: str, notes: str) -> str:
    """Heuristic category from title + notes. Used by both the ICS parser
    and the JSON parser."""
    combined = f"{title or ''} {notes or ''}".lower()
    if any(w in combined for w in ["flight", "airline", "airport", "terminal"]):
        return "flight"
    if any(w in combined for w in ["train", "rail", "tgv", "ave ", "eurostar", "amtrak"]):
        return "train"
    if any(w in combined for w in ["bus", "coach"]):
        return "bus"
    if any(w in combined for w in ["car rental", "uber", "taxi", "transfer"]):
        return "transport"
    if any(
        w in combined
        for w in ["hotel", "hostel", "airbnb", "accommodation", "check-in", "check in", "stay"]
    ):
        return "hotel"
    if any(
        w in combined
        for w in ["restaurant", "dinner", "lunch", "breakfast", "cafe", "brunch", "reservation"]
    ):
        return "meal"
    if any(w in combined for w in ["museum", "tour", "visit", "cathedral", "palace", "gallery"]):
        return "attraction"
    return "activity"


def _parse_ics_file(file_data: bytes) -> list[dict[str, Any]]:
    """Parse ICS calendar file to extract travel events.

    Uses the icalendar library so we get correct line folding, parameter
    parsing (TZID etc.), escape handling, and timezone-aware datetimes for
    free. The previous hand-rolled version ignored property params and
    folded lines with tabs, which truncated real-world calendar files.

    Returns list of items with title, category, date, time, location, notes.
    """
    from datetime import date as _date
    from datetime import datetime as _datetime

    from icalendar import Calendar

    try:
        cal = Calendar.from_ical(file_data)
    except Exception as e:
        # Malformed input. Caller treats this as "no items found" via the
        # empty-list contract; surface enough context so the upload handler
        # can show something useful.
        print(f"[ics-import] icalendar parse failed: {e}")
        return []

    items: list[dict[str, Any]] = []
    for component in cal.walk("VEVENT"):
        title = str(component.get("SUMMARY", "")).strip()
        if not title:
            continue

        item: dict[str, Any] = {"title": title}

        dtstart = component.get("DTSTART")
        if dtstart is not None:
            dt_value = dtstart.dt
            if isinstance(dt_value, _datetime):
                # Use the local wall-clock time of the event. If TZID was
                # set, dt_value is timezone-aware and .strftime gives us
                # the local time in that zone. If it was floating time, we
                # also just use it as-is. Converting to UTC would shift a
                # 4pm SFO flight to 11pm in the calendar app.
                item["date"] = dt_value.strftime("%Y-%m-%d")
                item["time"] = dt_value.strftime("%H:%M")
            elif isinstance(dt_value, _date):
                item["date"] = dt_value.strftime("%Y-%m-%d")

        dtend = component.get("DTEND")
        if dtend is not None:
            dt_end = dtend.dt
            if isinstance(dt_end, _datetime):
                item["end_time"] = dt_end.strftime("%H:%M")

        location = component.get("LOCATION")
        if location:
            item["location"] = str(location).strip()

        description = component.get("DESCRIPTION")
        if description:
            desc_text = str(description).strip()
            item["notes"] = desc_text[:500]

            # TripIt and several travel-booking exporters stamp DTSTART as
            # UTC and put the actual local departure time in the description.
            # Override the UTC-derived time with the local one when present.
            time_match = re.search(
                r"(?:Departure time|Departs?):\s*(\d{1,2}):(\d{2})",
                desc_text,
                re.IGNORECASE,
            )
            if time_match:
                hour = int(time_match.group(1))
                minute = time_match.group(2)
                item["time"] = f"{hour:02d}:{minute}"

        item["category"] = _guess_event_category(item.get("title", ""), item.get("notes", ""))
        items.append(item)

    return items


def _parse_json_trip(file_data: bytes) -> list[dict[str, Any]]:
    """Parse JSON file that might contain trip data.

    Handles various JSON formats:
    - Our own export format with itinerary_data
    - Our own itinerary format with items array
    - Array of events
    - TripIt-style JSON exports
    """
    try:
        content = file_data.decode("utf-8")
    except UnicodeDecodeError:
        content = file_data.decode("latin-1")

    data = json.loads(content)
    items = []

    if isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                items.append(_normalize_item(item))
    elif isinstance(data, dict):
        if "export_version" in data and "itinerary_data" in data:
            itinerary_data = data.get("itinerary_data", {})
            for day in itinerary_data.get("days", []):
                day_num = day.get("day_number") or day.get("day")
                day_date = day.get("date")
                for item in day.get("items", []):
                    normalized = _normalize_item(item)
                    if day_num and not normalized.get("day"):
                        normalized["day"] = day_num
                    if day_date and not normalized.get("date"):
                        normalized["date"] = day_date
                    items.append(normalized)
            for item in itinerary_data.get("ideas", []):
                items.append(_normalize_item(item))
        elif "items" in data:
            for item in data.get("items", []):
                if isinstance(item, dict):
                    items.append(_normalize_item(item))
        elif "days" in data:
            for day in data.get("days", []):
                day_num = day.get("day_number") or day.get("day")
                day_date = day.get("date")
                for item in day.get("items", []):
                    normalized = _normalize_item(item)
                    if day_num and not normalized.get("day"):
                        normalized["day"] = day_num
                    if day_date and not normalized.get("date"):
                        normalized["date"] = day_date
                    items.append(normalized)
        elif "itinerary_data" in data:
            return _parse_json_trip(json.dumps(data["itinerary_data"]).encode())
        elif "events" in data:
            for item in data.get("events", []):
                if isinstance(item, dict):
                    items.append(_normalize_item(item))

    # Smart day assignment: if no items have dates but some have times,
    # this is likely a single-day itinerary - assign all to Day 1
    has_any_date = any(item.get("date") for item in items)
    has_any_time = any(item.get("time") for item in items)
    has_any_day = any(item.get("day") for item in items)

    if not has_any_date and not has_any_day and has_any_time:
        for item in items:
            item["day"] = 1

    return items


def _parse_excel_to_text(file_data: bytes, ext: str) -> str:
    """Parse Excel file and convert to text table for LLM processing."""
    from io import BytesIO

    try:
        import openpyxl
    except ImportError as e:
        raise ImportError("openpyxl not installed") from e

    text_parts = []

    if ext == "xlsx":
        wb = openpyxl.load_workbook(BytesIO(file_data), data_only=True)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===\n")

            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            header_row = None
            for i, row in enumerate(rows):
                if any(cell is not None for cell in row):
                    header_row = i
                    break

            if header_row is None:
                continue

            for row in rows[header_row:]:
                row_text = [str(cell) if cell is not None else "" for cell in row]
                if any(row_text):
                    text_parts.append(" | ".join(row_text))

            text_parts.append("")

        wb.close()
    elif ext == "xls":
        try:
            import xlrd

            wb = xlrd.open_workbook(file_contents=file_data)

            for sheet_name in wb.sheet_names():
                sheet = wb.sheet_by_name(sheet_name)
                text_parts.append(f"=== Sheet: {sheet_name} ===\n")

                for row_idx in range(sheet.nrows):
                    row_text = [
                        str(sheet.cell_value(row_idx, col_idx)) or ""
                        for col_idx in range(sheet.ncols)
                    ]
                    if any(row_text):
                        text_parts.append(" | ".join(row_text))

                text_parts.append("")
        except ImportError as e:
            raise ImportError("xlrd not installed for .xls files") from e

    return "\n".join(text_parts)


def _parse_word_to_text(file_data: bytes, ext: str) -> str:
    """Parse Word document and extract text for LLM processing."""
    from io import BytesIO

    if ext == "docx":
        try:
            from docx import Document
        except ImportError as e:
            raise ImportError("python-docx not installed") from e

        doc = Document(BytesIO(file_data))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            text_parts.append("\n--- Table ---")
            headers: list[str] = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    # Treat first row as headers for context
                    headers = cells
                    text_parts.append(" | ".join(cells))
                elif headers:
                    # Annotate each cell with its column name for LLM clarity
                    annotated = [
                        f"{headers[i]}: {cell}"
                        if i < len(headers) and headers[i] and cell
                        else cell
                        for i, cell in enumerate(cells)
                    ]
                    text_parts.append(" | ".join(annotated))
                else:
                    text_parts.append(" | ".join(cells))

        return "\n".join(text_parts)

    elif ext == "doc":
        raise ValueError("Legacy .doc format not supported. Please save as .docx")
