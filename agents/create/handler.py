"""API handlers for the Create Trip agent."""

import json
import os
import re
import ssl
import urllib.request
from datetime import datetime
from typing import Any, Dict, List, Optional

from anthropic import Anthropic

import database as db


def _fetch_webpage_for_chat(url: str) -> dict:
    """Fetch a web page and extract text for chat handlers.

    Returns dict with success, text, title, error fields.
    """
    try:
        # Create SSL context
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        }

        req = urllib.request.Request(url, headers=headers)

        with urllib.request.urlopen(req, context=ctx, timeout=60) as response:
            content = response.read()

        # Extract text from HTML
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self.skip_tags = {'script', 'style', 'meta', 'link', 'noscript'}
                self.current_skip = False

            def handle_starttag(self, tag, attrs):
                if tag in self.skip_tags:
                    self.current_skip = True
                elif tag in ('br', 'p', 'div', 'li', 'tr', 'h1', 'h2', 'h3'):
                    self.text_parts.append('\n')

            def handle_endtag(self, tag):
                if tag in self.skip_tags:
                    self.current_skip = False

            def handle_data(self, data):
                if not self.current_skip:
                    text = data.strip()
                    if text:
                        self.text_parts.append(text + ' ')

        try:
            html_str = content.decode('utf-8')
        except UnicodeDecodeError:
            html_str = content.decode('latin-1')

        extractor = TextExtractor()
        extractor.feed(html_str)
        text = ''.join(extractor.text_parts)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text).strip()

        # Extract title
        title = None
        title_match = re.search(r'<title[^>]*>([^<]+)</title>', html_str, re.IGNORECASE)
        if title_match:
            title = title_match.group(1).strip()

        # Limit text length
        if len(text) > 15000:
            text = text[:15000] + "\n\n[Content truncated...]"

        return {
            'success': True,
            'text': text,
            'title': title or url,
            'url': url
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'url': url
        }


def _load_curated_venues() -> List[Dict]:
    """Load curated venues from database for cross-referencing."""
    try:
        venues = db.get_all_venues()
        return venues if venues else []
    except Exception as e:
        print(f"Error loading curated venues: {e}")
        return []


def _cross_reference_curated(name: str, venues: List[Dict]) -> Optional[Dict]:
    """Check if a venue name exists in the curated database."""
    name_lower = name.lower().strip()
    for v in venues:
        if v.get('name', '').lower() == name_lower:
            return v
        # Fuzzy match
        if name_lower in v.get('name', '').lower() or v.get('name', '').lower() in name_lower:
            return v
    return None


def create_trip_handler(user_id: int, data: Dict[str, Any]) -> Dict[str, Any]:
    """Create a new draft trip.

    Args:
        user_id: The user's ID
        data: Request data containing title, start_date, end_date, num_days

    Returns:
        The created trip data or error
    """
    title = data.get('title', '').strip()
    if not title:
        return {'error': 'Trip title is required'}, 400

    start_date = data.get('start_date')
    end_date = data.get('end_date')
    num_days = data.get('num_days')

    # Validate num_days
    if num_days is not None:
        try:
            num_days = int(num_days)
            if num_days < 1 or num_days > 365:
                return {'error': 'Number of days must be between 1 and 365'}, 400
        except (ValueError, TypeError):
            return {'error': 'Invalid number of days'}, 400

    trip = db.create_draft_trip(
        user_id=user_id,
        title=title,
        start_date=start_date,
        end_date=end_date,
        num_days=num_days
    )

    if trip:
        return {'success': True, 'trip': trip}, 200
    else:
        return {'error': 'Failed to create trip'}, 500


def save_trip_handler(user_id: int, link: str, data: Dict[str, Any]) -> Dict[str, Any]:
    """Auto-save trip itinerary data.

    Args:
        user_id: The user's ID
        link: The trip's unique link
        data: Request data containing itinerary_data

    Returns:
        Success or error response
    """
    itinerary_data = data.get('itinerary_data')
    if itinerary_data is None:
        return {'error': 'No itinerary data provided'}, 400

    success = db.update_trip_itinerary_data(user_id, link, itinerary_data)

    if success:
        # If trip is already published (not a draft), regenerate HTML
        trip = db.get_trip_by_link(user_id, link)
        if trip and not trip.get('is_draft', True):
            _generate_trip_html(trip, link)

        return {'success': True, 'saved_at': datetime.now().isoformat()}, 200
    else:
        return {'error': 'Failed to save trip'}, 500


def _generate_trip_html(trip: Dict[str, Any], link: str) -> bool:
    """Generate HTML file for a trip.

    Args:
        trip: Trip data from database
        link: The trip's link (filename)

    Returns:
        True if successful, False otherwise
    """
    import os
    from pathlib import Path

    try:
        itinerary = _convert_to_itinerary(trip)
        if itinerary and itinerary.items:
            from agents.itinerary.web_view import ItineraryWebView

            OUTPUT_DIR = Path(os.environ.get("OUTPUT_DIR", Path(__file__).parent.parent.parent / "output"))
            web_view = ItineraryWebView()
            web_view.generate(itinerary, OUTPUT_DIR / link, use_ai_summary=False, skip_geocoding=True)
            print(f"Generated HTML for trip: {link}")
            return True
    except Exception as e:
        print(f"Warning: Could not generate trip HTML: {e}")
        import traceback
        traceback.print_exc()

    return False


def publish_trip_handler(user_id: int, link: str) -> Dict[str, Any]:
    """Publish a draft trip (set is_draft=False) and generate HTML.

    Args:
        user_id: The user's ID
        link: The trip's unique link

    Returns:
        Success or error response
    """
    # Get the trip data first
    trip = db.get_trip_by_link(user_id, link)
    if not trip:
        return {'error': 'Trip not found'}, 404

    # Generate the HTML file from itinerary_data
    _generate_trip_html(trip, link)

    # Update database to mark as published
    success = db.publish_draft(user_id, link)

    if success:
        return {'success': True}, 200
    else:
        return {'error': 'Failed to publish trip'}, 500


def _convert_to_itinerary(trip: Dict[str, Any]):
    """Convert trip data from database to Itinerary object for HTML generation."""
    from datetime import datetime, date, time
    from agents.itinerary.models import Itinerary, ItineraryItem, Location

    itinerary_data = trip.get('itinerary_data') or {}

    # Get title
    title = itinerary_data.get('title') or trip.get('title', 'Untitled Trip')

    # Collect all items from days and ideas
    items = []

    # Process days
    days = itinerary_data.get('days', [])
    for day in days:
        day_number = day.get('day_number')
        day_date_str = day.get('date')
        day_date = None
        if day_date_str:
            try:
                day_date = datetime.strptime(day_date_str, '%Y-%m-%d').date()
            except:
                pass

        for item_data in day.get('items', []):
            item = _create_itinerary_item(item_data, day_number, day_date)
            if item:
                items.append(item)

    # Process ideas pile (items without dates)
    ideas = itinerary_data.get('ideas', [])
    for item_data in ideas:
        item = _create_itinerary_item(item_data, None, None)
        if item:
            items.append(item)

    # Get start/end dates
    start_date = None
    end_date = None
    if itinerary_data.get('start_date'):
        try:
            start_date = datetime.strptime(itinerary_data['start_date'], '%Y-%m-%d').date()
        except:
            pass
    if itinerary_data.get('end_date'):
        try:
            end_date = datetime.strptime(itinerary_data['end_date'], '%Y-%m-%d').date()
        except:
            pass

    return Itinerary(
        title=title,
        items=items,
        start_date=start_date,
        end_date=end_date,
        travelers=itinerary_data.get('travelers', []),
        source_file=None
    )


def _create_itinerary_item(item_data: Dict[str, Any], day_number: Optional[int], day_date) -> Optional:
    """Create an ItineraryItem from create trip item data."""
    from datetime import datetime, time
    from agents.itinerary.models import ItineraryItem, Location

    if not item_data.get('title'):
        return None

    # Parse location
    location_data = item_data.get('location', '')
    if isinstance(location_data, dict):
        location_name = location_data.get('name', '') or location_data.get('city', '')
    else:
        location_name = str(location_data) if location_data else ''

    location = Location(
        name=location_name,
        address=None,
        location_type=item_data.get('category')
    )

    # Parse start time
    start_time = None
    time_str = item_data.get('time')
    if time_str and isinstance(time_str, str) and ':' in time_str:
        try:
            parts = time_str.split(':')
            start_time = time(int(parts[0]), int(parts[1]))
        except:
            pass

    # Parse end time (for flights, trains, etc.)
    end_time_obj = None
    end_time_str = item_data.get('end_time')
    if end_time_str and isinstance(end_time_str, str) and ':' in end_time_str:
        try:
            parts = end_time_str.split(':')
            end_time_obj = time(int(parts[0]), int(parts[1]))
        except:
            pass

    # Map category
    category = item_data.get('category', 'activity')

    return ItineraryItem(
        title=item_data.get('title', 'Untitled'),
        location=location,
        date=day_date,
        start_time=start_time,
        end_time=end_time_obj,
        description=item_data.get('notes'),
        category=category,
        confirmation_number=None,
        notes=item_data.get('notes'),
        day_number=day_number,
        is_home_location=False,
        website_url=item_data.get('website')
    )


def get_trip_data_handler(user_id: int, link: str) -> Dict[str, Any]:
    """Get trip data for editing.

    Args:
        user_id: The user's ID
        link: The trip's unique link

    Returns:
        Trip data or error
    """
    trip = db.get_trip_by_link(user_id, link)

    if trip:
        return {'success': True, 'trip': trip}, 200
    else:
        return {'error': 'Trip not found'}, 404


def export_trip_handler(user_id: int, link: str) -> Dict[str, Any]:
    """Export trip data as downloadable JSON.

    Args:
        user_id: The user's ID
        link: The trip's unique link

    Returns:
        Trip data for export or error
    """
    trip = db.get_trip_by_link(user_id, link)

    if not trip:
        return {'error': 'Trip not found'}, 404

    # Build export data with all relevant fields
    export_data = {
        'export_version': '1.0',
        'exported_at': datetime.now().isoformat(),
        'title': trip.get('title', 'Untitled Trip'),
        'dates': trip.get('dates'),
        'days': trip.get('days'),
        'locations': trip.get('locations'),
        'activities': trip.get('activities'),
        'itinerary_data': trip.get('itinerary_data'),
        'is_public': trip.get('is_public', False),
    }

    return {'success': True, 'export': export_data}, 200


def add_item_to_trip_handler(user_id: int, link: str, data: Dict[str, Any]) -> Dict[str, Any]:
    """Add an item to trip's ideas pile.

    Args:
        user_id: The user's ID
        link: The trip's unique link
        data: Request data containing the item

    Returns:
        Success or error response
    """
    item = data.get('item')
    if not item:
        return {'error': 'No item provided'}, 400

    # Ensure item has required fields
    if 'title' not in item:
        return {'error': 'Item must have a title'}, 400

    success = db.add_item_to_trip(user_id, link, item)

    if success:
        return {'success': True}, 200
    else:
        return {'error': 'Failed to add item to trip'}, 500


def create_chat_handler(user_id: int, data: Dict[str, Any]) -> Dict[str, Any]:
    """Handle LLM chat for venue recommendations.

    Supports:
    - Curated venue database cross-referencing
    - Web page fetching for external lists (Eater, etc.)
    - Source tagging (CURATED vs AI_PICK)

    Args:
        user_id: The user's ID
        data: Request data containing message and history

    Returns:
        LLM response with suggested items
    """
    message = data.get('message', '').strip()
    if not message:
        return {'error': 'No message provided'}, 400

    history = data.get('history', [])
    trip_context = data.get('trip_context', {})

    # Load curated venues for cross-referencing
    curated_venues = _load_curated_venues()

    # Build system prompt for venue-focused chat
    system_prompt = _build_venue_chat_prompt(trip_context, curated_venues)

    # Build messages for LLM (filter out empty messages)
    messages = []
    for msg in history[-10:]:  # Last 10 messages for context
        content = msg.get('content', '').strip()
        if content:  # Only include messages with non-empty content
            messages.append({
                'role': msg.get('role', 'user'),
                'content': content
            })
    messages.append({'role': 'user', 'content': message})

    # Define tools for trip building and web fetching
    tools = [
        {
            "name": "add_to_itinerary",
            "description": "Add one or more items to the user's trip itinerary. Use this tool whenever the user asks to add, include, schedule, book, or plan something for their trip.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "items": {
                        "type": "array",
                        "description": "List of items to add to the trip",
                        "items": {
                            "type": "object",
                            "properties": {
                                "title": {
                                    "type": "string",
                                    "description": "Name of the place or activity"
                                },
                                "category": {
                                    "type": "string",
                                    "enum": ["meal", "hotel", "activity", "attraction", "transport", "other"],
                                    "description": "Type of item"
                                },
                                "location": {
                                    "type": "string",
                                    "description": "City or address"
                                },
                                "notes": {
                                    "type": "string",
                                    "description": "Additional details about the item"
                                },
                                "day": {
                                    "type": "integer",
                                    "description": "Day number to add to (1, 2, 3...). Omit to add to Ideas pile."
                                },
                                "time": {
                                    "type": "string",
                                    "description": "Time in 24-hour format like '14:30' (optional)"
                                },
                                "website": {
                                    "type": "string",
                                    "description": "Official website URL for the place (e.g., https://example.com)"
                                },
                                "source": {
                                    "type": "string",
                                    "enum": ["CURATED", "AI_PICK"],
                                    "description": "CURATED if from the venue database, AI_PICK if a new recommendation"
                                }
                            },
                            "required": ["title", "category"]
                        }
                    }
                },
                "required": ["items"]
            }
        },
        {
            "name": "fetch_web_page",
            "description": "Fetch a web page to extract venue recommendations. Use this when users mention external lists like Eater, Infatuation, blog posts, or provide URLs.",
            "input_schema": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The URL to fetch. For 'Eater 38 Rome', try 'https://www.eater.com/maps/best-restaurants-rome'"
                    }
                },
                "required": ["url"]
            }
        }
    ]

    try:
        client = Anthropic()

        # Tool use loop - handle multiple rounds if Claude calls tools
        max_iterations = 3
        web_fetch_context = None
        add_items = []
        response_text = ""

        for iteration in range(max_iterations):
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2048,
                system=system_prompt,
                messages=messages,
                tools=tools
            )

            # Process response blocks
            tool_use_block = None

            for block in response.content:
                if block.type == "text":
                    response_text = block.text
                elif block.type == "tool_use":
                    if block.name == "add_to_itinerary":
                        tool_input = block.input
                        if "items" in tool_input:
                            add_items = tool_input["items"]
                    elif block.name == "fetch_web_page":
                        tool_use_block = block

            # Handle web fetch tool
            if tool_use_block and tool_use_block.name == "fetch_web_page":
                url = tool_use_block.input.get("url", "")
                print(f"[CREATE CHAT] Fetching web page: {url}")

                fetch_result = _fetch_webpage_for_chat(url)

                if fetch_result['success']:
                    web_fetch_context = {
                        'url': url,
                        'title': fetch_result.get('title', url)
                    }
                    tool_result_content = f"Successfully fetched page: {fetch_result['title']}\n\nContent:\n{fetch_result['text']}"
                else:
                    tool_result_content = f"Failed to fetch page: {fetch_result.get('error', 'Unknown error')}"

                # Add assistant message with tool use and tool result
                messages.append({"role": "assistant", "content": response.content})
                messages.append({
                    "role": "user",
                    "content": [{
                        "type": "tool_result",
                        "tool_use_id": tool_use_block.id,
                        "content": tool_result_content
                    }]
                })
                # Continue loop for Claude to process the result
                continue

            # No web fetch - we have the final response
            break

        # Clean response text (remove any leftover JSON block for display)
        display_text = _clean_response_text(response_text)

        # Parse suggested items and cross-reference with curated DB
        suggested_items = _parse_suggested_items(display_text, curated_venues, web_fetch_context)

        # Add source field to add_items too
        for item in add_items:
            if 'source' not in item:
                curated_match = _cross_reference_curated(item.get('title', ''), curated_venues)
                item['source'] = 'CURATED' if curated_match else 'AI_PICK'
            if web_fetch_context and not item.get('collection'):
                item['collection'] = web_fetch_context.get('title', '')[:50]

        # Debug logging
        print(f"[CREATE CHAT] Response text length: {len(display_text)}")
        print(f"[CREATE CHAT] Parsed suggested items: {len(suggested_items)}")
        for i, item in enumerate(suggested_items[:5]):  # Show first 5 items
            source = item.get('source', 'NONE')
            print(f"[CREATE CHAT] Item {i+1}: {item.get('title', 'NO TITLE')} - source={source}")

        return {
            'success': True,
            'response': display_text,
            'suggested_items': suggested_items,
            'add_items': add_items
        }, 200

    except Exception as e:
        print(f"Create chat error: {e}")
        import traceback
        traceback.print_exc()
        return {'error': f'Chat service error: {str(e)}'}, 500


def _build_venue_chat_prompt(trip_context: Dict[str, Any], curated_venues: List[Dict] = None) -> str:
    """Build the system prompt for venue-focused chat."""

    destination = trip_context.get('destination', 'your destination')
    dates = trip_context.get('dates', '')
    days = trip_context.get('days', [])
    ideas = trip_context.get('ideas', [])
    curated_venues = curated_venues or []

    # Build context about what's already planned
    itinerary_context = ""

    if days:
        itinerary_context += "\n\nCurrent itinerary by day:"
        for day in days:
            day_num = day.get('day_number', '?')
            day_date = day.get('date', 'TBD')
            items = day.get('items', [])
            if items:
                itinerary_context += f"\n  Day {day_num} ({day_date}):"
                for item in items:
                    title = item.get('title', 'Untitled')
                    category = item.get('category', '')
                    time = item.get('time', '')
                    location = item.get('location', '')
                    time_str = f" at {time}" if time else ""
                    loc_str = f" - {location}" if location else ""
                    itinerary_context += f"\n    - [{category}] {title}{time_str}{loc_str}"
            else:
                itinerary_context += f"\n  Day {day_num} ({day_date}): No activities planned yet"

    if ideas:
        itinerary_context += "\n\nIdeas pile (items user is considering but hasn't scheduled):"
        for item in ideas:
            title = item.get('title', 'Untitled')
            category = item.get('category', '')
            notes = item.get('notes', '')
            notes_str = f" - {notes[:50]}..." if notes and len(notes) > 50 else (f" - {notes}" if notes else "")
            itinerary_context += f"\n  - [{category}] {title}{notes_str}"

    # Build day reference for adding to specific days
    day_reference = ""
    if days:
        day_reference = "\n\nAvailable days to add items to:"
        for day in days:
            day_num = day.get('day_number', '?')
            day_date = day.get('date', 'TBD')
            day_reference += f"\n  - Day {day_num} ({day_date})"

    # Build list of existing item titles for deduplication
    existing_titles = set()
    for day in days:
        for item in day.get('items', []):
            title = item.get('title', '').lower().strip()
            if title:
                existing_titles.add(title)
    for item in ideas:
        title = item.get('title', '').lower().strip()
        if title:
            existing_titles.add(title)

    existing_list = "\n".join(f"  - {t}" for t in sorted(existing_titles)) if existing_titles else "  (none yet)"

    # Build curated venue summary for context
    curated_context = ""
    if curated_venues:
        # Filter venues relevant to destination
        dest_lower = destination.lower() if destination else ''
        relevant_venues = [v for v in curated_venues
                          if dest_lower in (v.get('city', '') or '').lower()
                          or dest_lower in (v.get('country', '') or '').lower()
                          or dest_lower in (v.get('state', '') or '').lower()]

        if relevant_venues:
            curated_context = f"\n\n## CURATED VENUES DATABASE\n\nYou have access to {len(curated_venues)} vetted venues. Here are {len(relevant_venues)} venues near {destination}:\n"
            for v in relevant_venues[:50]:  # Limit to 50 relevant venues
                curated_context += f"- {v['name']}"
                if v.get('city'):
                    curated_context += f", {v['city']}"
                if v.get('venue_type'):
                    curated_context += f" ({v['venue_type']})"
                if v.get('michelin_stars'):
                    curated_context += f" ⭐{v['michelin_stars']} Michelin"
                if v.get('collection') and v['collection'] not in ('Saved', None):
                    curated_context += f" #{v['collection']}"
                curated_context += "\n"
            curated_context += "\nVenues from this list should be marked as source: CURATED\n"

    prompt = f"""You are a helpful travel planning assistant for Libertas, a travel itinerary app.

You have the ability to:
1. Add items to the user's itinerary using the add_to_itinerary tool
2. Fetch web pages using the fetch_web_page tool for external lists (Eater, Infatuation, blogs)

Current trip context:
- Destination: {destination}
- Dates: {dates if dates else 'Not set yet'}
{itinerary_context}
{day_reference}
{curated_context}

## CRITICAL: AVOID DUPLICATES

The user already has these items in their trip (DO NOT suggest or add these again):
{existing_list}

When the user asks for "other", "more", "different", or "alternative" options, you MUST suggest DIFFERENT places than those listed above.

## WEB FETCH

Use the fetch_web_page tool when users mention:
- External lists: "Eater 38", "Infatuation", "Michelin Guide", blog posts
- Specific URLs they want to check
- "Check this page for recommendations"

## WHEN TO USE add_to_itinerary TOOL

ONLY use the tool when the user EXPLICITLY asks to ADD something:
- "add this", "add these", "put this in my trip", "include this", "book this"

DO NOT use the tool for: "suggestions", "recommend", "ideas", "options", "what are some"

When using add_to_itinerary, include the source field:
- source: "CURATED" - if the venue is in the curated database above
- source: "AI_PICK" - if it's a new recommendation not in the database

Categories: meal, hotel, activity, attraction, transport, other
Day: Use day number (1, 2, 3...) or omit to add to Ideas pile

## SPECIFIC PLACE REQUESTS

When the user asks about a SPECIFIC place by name (e.g., "ABBA Museum", "Eiffel Tower", "Noma"):

REQUIRED FORMAT - always include this line:
**Venue Name** - Brief description of the place.

Example for "ABBA Museum":
**ABBA Museum** - Interactive museum on Djurgården celebrating Sweden's legendary pop group with costumes, memorabilia, and singalong experiences.

Then optionally add a plain text tip (no bold). Do NOT list features or use bullet points.

## GENERAL SUGGESTIONS (only when asked)

When the user asks for general suggestions ("recommend restaurants", "what should I see", "ideas for activities"):
- Provide 3-5 options in a numbered list
- Format: **Venue Name** - Brief description. [Website](url)
- Do NOT use bold text (**) for anything except venue names

## FORMATTING RULES

CRITICAL: Only use **bold** for venue/place names that can be added to the itinerary.

NEVER use bold for:
- Features, highlights, exhibits, menu items
- Options like "Get more information" or "Add to itinerary"
- Questions or action choices
- Dates, times, or day references

Bad examples (creates multiple unwanted suggestions):
- **ABBA Museum** with **Costume exhibits** and **Audio guide**
- What would you prefer? **Get more info** or **Add to itinerary**?

Good example (creates exactly one suggestion):
**ABBA Museum** - Interactive museum celebrating Sweden's famous pop group, featuring costumes, holograms, and singalong booths. Book tickets in advance online.

Want me to add it to a specific day?
"""
    return prompt


def _parse_add_items(response_text: str) -> List[Dict[str, Any]]:
    """Parse items to add from JSON block in LLM response.

    Looks for ```json blocks with add_items array.
    Returns list of items to directly add to the trip.
    """
    import re

    # Look for JSON code block with add_items
    json_pattern = r'```json\s*(\{[\s\S]*?"add_items"[\s\S]*?\})\s*```'
    match = re.search(json_pattern, response_text)

    if match:
        try:
            data = json.loads(match.group(1))
            if 'add_items' in data and isinstance(data['add_items'], list):
                return data['add_items']
        except json.JSONDecodeError as e:
            print(f"Failed to parse add_items JSON: {e}")

    return []


def _clean_response_text(response_text: str) -> str:
    """Remove the JSON block from the response text for display."""
    import re
    # Remove the JSON code block
    cleaned = re.sub(r'```json\s*\{[\s\S]*?"add_items"[\s\S]*?\}\s*```', '', response_text)
    return cleaned.strip()


def _parse_suggested_items(response_text: str, curated_venues: List[Dict] = None, web_fetch_context: Dict = None) -> List[Dict[str, Any]]:
    """Parse suggested items from the LLM response.

    This is a simple parser that looks for structured suggestions.
    Cross-references with curated database to add source field.
    Returns a list of items that can be added to the trip.
    """
    items = []
    curated_venues = curated_venues or []

    # Pattern 1: Numbered items with bold names like "1. **Name** - description"
    pattern1 = r'\d+\.\s+\*\*([^*]+)\*\*\s*[-–—:]?\s*(.+?)(?=\n\d+\.|\n\n|$)'
    matches = re.findall(pattern1, response_text, re.DOTALL)

    # Pattern 2: Bullet points with bold names like "- **Name** - description"
    if not matches:
        pattern2 = r'[-•]\s+\*\*([^*]+)\*\*\s*[-–—:]?\s*(.+?)(?=\n[-•]|\n\n|$)'
        matches = re.findall(pattern2, response_text, re.DOTALL)

    # Pattern 3: Just bold names on their own line "**Name**"
    if not matches:
        pattern3 = r'\*\*([^*]+)\*\*\s*[-–—:]?\s*([^\n*]+)?'
        matches = re.findall(pattern3, response_text)

    # Pattern 4: Plain text "Name - description" on separate lines (no bold)
    if not matches:
        # Look for lines that start with a capitalized word followed by " - " and description
        pattern4 = r'^([A-Z][A-Za-z\s&\']+?)\s*[-–—]\s*(.+?)$'
        matches = re.findall(pattern4, response_text, re.MULTILINE)
        # Filter out lines that are too long (likely regular sentences, not names)
        matches = [(m[0].strip(), m[1].strip()) for m in matches if len(m[0].strip()) < 50]

    for match in matches:
        name = match[0].strip() if match[0] else ''
        description = match[1].strip() if len(match) > 1 and match[1] else ''

        if not name:
            continue

        # Skip question-style items (Claude's follow-up questions, not actual venues)
        name_lower = name.lower()
        skip_phrases = [
            # Questions and actions
            'want me to', 'would you like', 'shall i', 'should i', 'let me know',
            'add it to', 'get more', 'suggest other', 'nearby', 'something else',
            'more information', 'what you', 'i can', 'i already', 'i shared',
            # Itinerary/trip references
            'your itinerary', 'your trip', 'which day', 'if so', 'available days',
            'day 1', 'day 2', 'day 3', 'day 4', 'day 5',
            # Dates
            'dec ', 'december', 'january', 'february', 'march', 'april', 'may ',
            'june', 'july', 'august', 'september', 'october', 'november',
            # Generic options
            'option', 'prefer', 'choose', 'select', 'pick one',
        ]
        if any(q in name_lower for q in skip_phrases):
            continue
        # Skip if name ends with question mark or colon (headings)
        if name.rstrip().endswith('?') or name.rstrip().endswith(':'):
            continue
        # Skip generic/vague names
        skip_exact = ['yes', 'no', 'here', 'there', 'this', 'that', 'more', 'other',
                      'something else', 'get more information', 'add it to your itinerary']
        if name_lower.strip() in skip_exact:
            continue
        # Skip names that are too long (likely sentences, not venue names)
        if len(name) > 60:
            continue
        # Skip names that start with common non-venue words
        if name_lower.startswith(('i ', 'you ', 'we ', 'let ', 'if ', 'what ', 'how ', 'why ', 'when ', 'where ')):
            continue

        # Extract website URL from markdown format [text](url) or plain URL
        website = None

        # Pattern 1: Any markdown link format [any text](url)
        any_link_pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
        url_match = re.search(any_link_pattern, description)
        if url_match:
            website = url_match.group(2)  # Group 2 is the URL
            # Remove the markdown link from description
            description = re.sub(any_link_pattern, '', description).strip()
        else:
            # Pattern 2: Plain URL anywhere in the text (not just at end)
            plain_url = r'(https?://[^\s\)\]]+)'
            plain_match = re.search(plain_url, description)
            if plain_match:
                website = plain_match.group(1)
                description = re.sub(plain_url, '', description).strip()

        # Clean up trailing punctuation from description
        description = description.rstrip(' .-–—')

        # Try to determine category from keywords
        category = 'activity'
        combined_lower = (name + ' ' + description).lower()
        if any(word in combined_lower for word in ['restaurant', 'cafe', 'bakery', 'deli', 'trattoria', 'food', 'cuisine', 'dishes', 'dining', 'bistro']):
            category = 'meal'
        elif any(word in combined_lower for word in ['hotel', 'hostel', 'stay', 'accommodation', 'rooms', 'inn', 'lodge']):
            category = 'hotel'
        elif any(word in combined_lower for word in ['museum', 'gallery', 'cathedral', 'church', 'monument', 'palace', 'castle', 'theater', 'theatre', 'opera', 'concert hall']):
            category = 'attraction'
        elif any(word in combined_lower for word in ['hike', 'trail', 'tour', 'trek', 'walk', 'cycling']):
            category = 'activity'

        # Cross-reference with curated database
        curated_match = _cross_reference_curated(name, curated_venues)
        source = 'CURATED' if curated_match else 'AI_PICK'

        item = {
            'title': name,
            'category': category,
            'notes': description[:200] if len(description) > 200 else description,
            'source': source
        }

        # Add collection from curated match or web fetch context
        if curated_match and curated_match.get('collection'):
            item['collection'] = curated_match['collection']
        elif web_fetch_context:
            item['collection'] = web_fetch_context.get('title', '')[:50]

        if website:
            item['website'] = website

        # Add extra info from curated match
        if curated_match:
            if curated_match.get('website') and not website:
                item['website'] = curated_match['website']
            if curated_match.get('city'):
                item['location'] = curated_match['city']

        items.append(item)

    # Sort: CURATED first, then AI_PICK
    items.sort(key=lambda x: (0 if x.get('source') == 'CURATED' else 1, x.get('title', '')))

    return items


def _parse_ics_file(file_data: bytes) -> List[Dict[str, Any]]:
    """Parse ICS calendar file to extract travel events.

    Returns list of items with title, category, date, time, location, notes.
    """
    try:
        content = file_data.decode('utf-8')
    except UnicodeDecodeError:
        content = file_data.decode('latin-1')

    items = []
    current_event = {}
    in_event = False

    lines = content.replace('\r\n ', '').replace('\r\n\t', '').split('\r\n')
    if len(lines) == 1:
        lines = content.replace('\n ', '').replace('\n\t', '').split('\n')

    for line in lines:
        line = line.strip()
        if line == 'BEGIN:VEVENT':
            in_event = True
            current_event = {}
        elif line == 'END:VEVENT':
            if current_event.get('title'):
                # Determine category from summary/description
                title = current_event.get('title', '').lower()
                description = current_event.get('notes', '').lower()
                combined = f"{title} {description}"

                category = 'activity'
                if any(w in combined for w in ['flight', 'airline', 'airport', 'terminal']):
                    category = 'flight'
                elif any(w in combined for w in ['train', 'bus', 'car rental', 'uber', 'taxi', 'transfer']):
                    category = 'transport'
                elif any(w in combined for w in ['hotel', 'hostel', 'airbnb', 'accommodation', 'check-in', 'check in', 'stay']):
                    category = 'hotel'
                elif any(w in combined for w in ['restaurant', 'dinner', 'lunch', 'breakfast', 'cafe', 'brunch', 'reservation']):
                    category = 'meal'
                elif any(w in combined for w in ['museum', 'tour', 'visit', 'cathedral', 'palace', 'gallery']):
                    category = 'attraction'

                current_event['category'] = category

                # Use UTC time as fallback if no local time was extracted from description
                if not current_event.get('time') and current_event.get('_utc_time'):
                    current_event['time'] = current_event['_utc_time']

                # Clean up internal field
                current_event.pop('_utc_time', None)

                items.append(current_event)
            in_event = False
        elif in_event:
            if line.startswith('SUMMARY:'):
                current_event['title'] = line[8:].strip()
            elif line.startswith('DTSTART'):
                # Parse date/time - store raw value, may be overridden by description
                value = line.split(':', 1)[-1]
                if 'T' in value:
                    # Has time component
                    date_part = value[:8]
                    time_part = value[9:13] if len(value) > 12 else None
                    try:
                        current_event['date'] = f"{date_part[:4]}-{date_part[4:6]}-{date_part[6:8]}"
                        if time_part:
                            # Store UTC time as fallback - may be overridden by local time from description
                            current_event['_utc_time'] = f"{time_part[:2]}:{time_part[2:4]}"
                    except:
                        pass
                else:
                    # Date only
                    try:
                        current_event['date'] = f"{value[:4]}-{value[4:6]}-{value[6:8]}"
                    except:
                        pass
            elif line.startswith('LOCATION:'):
                current_event['location'] = line[9:].strip()
            elif line.startswith('DESCRIPTION:'):
                desc = line[12:].strip()
                # Unescape ICS format
                desc = desc.replace('\\n', '\n').replace('\\,', ',').replace('\\;', ';')
                current_event['notes'] = desc[:500]  # Limit notes length

                # Extract local departure time from description (more accurate than UTC)
                # Look for patterns like "Departure time: 10:15" or "Departs: 10:15"
                import re
                time_match = re.search(r'(?:Departure time|Departs?):\s*(\d{1,2}):(\d{2})', desc, re.IGNORECASE)
                if time_match:
                    hour = int(time_match.group(1))
                    minute = time_match.group(2)
                    current_event['time'] = f"{hour:02d}:{minute}"

    return items


def _parse_json_trip(file_data: bytes) -> List[Dict[str, Any]]:
    """Parse JSON file that might contain trip data.

    Handles various JSON formats:
    - Our own export format with itinerary_data
    - Our own itinerary format with items array
    - Array of events
    - TripIt-style JSON exports
    """
    try:
        content = file_data.decode('utf-8')
    except UnicodeDecodeError:
        content = file_data.decode('latin-1')

    data = json.loads(content)
    items = []

    # Handle different JSON structures
    if isinstance(data, list):
        # Direct array of items
        for item in data:
            if isinstance(item, dict):
                items.append(_normalize_item(item))
    elif isinstance(data, dict):
        # Check for our export format (has export_version and itinerary_data)
        if 'export_version' in data and 'itinerary_data' in data:
            itinerary_data = data.get('itinerary_data', {})
            # Process days from itinerary_data
            for day in itinerary_data.get('days', []):
                day_num = day.get('day_number') or day.get('day')
                day_date = day.get('date')
                for item in day.get('items', []):
                    normalized = _normalize_item(item)
                    if day_num and not normalized.get('day'):
                        normalized['day'] = day_num
                    if day_date and not normalized.get('date'):
                        normalized['date'] = day_date
                    items.append(normalized)
            # Process ideas pile
            for item in itinerary_data.get('ideas', []):
                items.append(_normalize_item(item))
        # Check for our itinerary format
        elif 'items' in data:
            for item in data.get('items', []):
                if isinstance(item, dict):
                    items.append(_normalize_item(item))
        # Check for days array
        elif 'days' in data:
            for day in data.get('days', []):
                day_num = day.get('day_number') or day.get('day')
                day_date = day.get('date')
                for item in day.get('items', []):
                    normalized = _normalize_item(item)
                    if day_num and not normalized.get('day'):
                        normalized['day'] = day_num
                    if day_date and not normalized.get('date'):
                        normalized['date'] = day_date
                    items.append(normalized)
        # Check for itinerary_data without export_version (just in case)
        elif 'itinerary_data' in data:
            return _parse_json_trip(json.dumps(data['itinerary_data']).encode())
        # Check for events array (common export format)
        elif 'events' in data:
            for item in data.get('events', []):
                if isinstance(item, dict):
                    items.append(_normalize_item(item))

    # Smart day assignment: if no items have dates but some have times,
    # this is likely a single-day itinerary - assign all to Day 1
    has_any_date = any(item.get('date') for item in items)
    has_any_time = any(item.get('time') for item in items)
    has_any_day = any(item.get('day') for item in items)

    if not has_any_date and not has_any_day and has_any_time:
        # Looks like a single-day itinerary without explicit dates
        # Assign all items to Day 1
        for item in items:
            item['day'] = 1

    return items


def _normalize_item(item: Dict[str, Any]) -> Dict[str, Any]:
    """Normalize an item dict to our expected format."""
    normalized = {}

    # Title can be in various fields
    normalized['title'] = (
        item.get('title') or
        item.get('name') or
        item.get('summary') or
        item.get('event') or
        'Untitled'
    )

    # Category normalization
    cat = (item.get('category') or item.get('type') or '').lower()
    if cat in ['flight', 'air', 'plane']:
        normalized['category'] = 'flight'
    elif cat in ['train', 'bus', 'car', 'transport', 'transportation', 'transfer']:
        normalized['category'] = 'transport'
    elif cat in ['hotel', 'accommodation', 'lodging', 'stay', 'hostel']:
        normalized['category'] = 'hotel'
    elif cat in ['meal', 'restaurant', 'food', 'dining', 'breakfast', 'lunch', 'dinner']:
        normalized['category'] = 'meal'
    elif cat in ['attraction', 'sightseeing', 'museum', 'tour']:
        normalized['category'] = 'attraction'
    elif cat in ['activity', 'event']:
        normalized['category'] = 'activity'
    else:
        normalized['category'] = cat or 'activity'

    # Date handling
    date = item.get('date') or item.get('start_date') or item.get('startDate')
    if date:
        # Try to parse various date formats
        if isinstance(date, str):
            # Already in ISO format
            if len(date) >= 10 and date[4] == '-':
                normalized['date'] = date[:10]
            # Try common formats
            else:
                for fmt in ['%Y-%m-%d', '%m/%d/%Y', '%d/%m/%Y', '%B %d, %Y']:
                    try:
                        parsed = datetime.strptime(date[:10], fmt)
                        normalized['date'] = parsed.strftime('%Y-%m-%d')
                        break
                    except:
                        pass

    # Time handling
    time = item.get('time') or item.get('start_time') or item.get('startTime')
    if time:
        if isinstance(time, str):
            # Handle HH:MM format
            if ':' in time and len(time) >= 5:
                normalized['time'] = time[:5]
            # Handle HHMM format
            elif len(time) == 4 and time.isdigit():
                normalized['time'] = f"{time[:2]}:{time[2:]}"

    # End time handling (for flights, trains, etc.)
    end_time = item.get('end_time') or item.get('endTime') or item.get('arrival_time')
    if end_time:
        if isinstance(end_time, str):
            # Handle HH:MM format
            if ':' in end_time and len(end_time) >= 5:
                normalized['end_time'] = end_time[:5]
            # Handle HHMM format
            elif len(end_time) == 4 and end_time.isdigit():
                normalized['end_time'] = f"{end_time[:2]}:{end_time[2:]}"

    # Location
    loc = item.get('location')
    if isinstance(loc, dict):
        normalized['location'] = loc.get('name') or loc.get('address') or loc.get('city')
    elif isinstance(loc, str):
        normalized['location'] = loc
    elif item.get('city'):
        normalized['location'] = item.get('city')
    elif item.get('address'):
        normalized['location'] = item.get('address')

    # Notes
    notes = item.get('notes') or item.get('description') or item.get('details')
    if notes:
        normalized['notes'] = str(notes)[:500]

    # Day number if present
    if item.get('day') or item.get('day_number'):
        normalized['day'] = item.get('day') or item.get('day_number')

    return normalized


def _parse_excel_to_text(file_data: bytes, ext: str) -> str:
    """Parse Excel file and convert to text table for LLM processing.

    Returns a text representation of the spreadsheet.
    """
    from io import BytesIO

    try:
        import openpyxl
    except ImportError:
        raise ImportError("openpyxl not installed")

    text_parts = []

    if ext == 'xlsx':
        wb = openpyxl.load_workbook(BytesIO(file_data), data_only=True)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===\n")

            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            # Find the header row (first non-empty row)
            header_row = None
            for i, row in enumerate(rows):
                if any(cell is not None for cell in row):
                    header_row = i
                    break

            if header_row is None:
                continue

            # Convert to text table
            for row in rows[header_row:]:
                row_text = []
                for cell in row:
                    if cell is not None:
                        row_text.append(str(cell))
                    else:
                        row_text.append('')
                if any(row_text):  # Skip empty rows
                    text_parts.append(' | '.join(row_text))

            text_parts.append('')  # Blank line between sheets

        wb.close()
    elif ext == 'xls':
        # For .xls files, try xlrd
        try:
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_data)

            for sheet_name in wb.sheet_names():
                sheet = wb.sheet_by_name(sheet_name)
                text_parts.append(f"=== Sheet: {sheet_name} ===\n")

                for row_idx in range(sheet.nrows):
                    row_text = []
                    for col_idx in range(sheet.ncols):
                        cell = sheet.cell_value(row_idx, col_idx)
                        row_text.append(str(cell) if cell else '')
                    if any(row_text):
                        text_parts.append(' | '.join(row_text))

                text_parts.append('')
        except ImportError:
            # Fall back to openpyxl's xlrd compatibility
            raise ImportError("xlrd not installed for .xls files")

    return '\n'.join(text_parts)


def _parse_word_to_text(file_data: bytes, ext: str) -> str:
    """Parse Word document and extract text for LLM processing.

    Returns the text content of the document.
    """
    from io import BytesIO

    if ext == 'docx':
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed")

        doc = Document(BytesIO(file_data))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract tables
        for table in doc.tables:
            text_parts.append('\n--- Table ---')
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text_parts.append(' | '.join(row_text))

        return '\n'.join(text_parts)

    elif ext == 'doc':
        # .doc files are harder - try antiword or textract
        # For now, return error suggesting conversion
        raise ValueError("Legacy .doc format not supported. Please save as .docx")


def upload_plan_handler(user_id: int, filename: str, file_data: bytes, ext: str) -> Dict[str, Any]:
    """Handle uploaded file and extract trip items using LLM.

    Args:
        user_id: The user's ID
        filename: Original filename
        file_data: Raw file bytes
        ext: File extension

    Returns:
        Extracted items or error
    """
    import base64
    from io import BytesIO

    # Determine how to process the file
    content_for_llm = None
    image_data = None
    media_type = None
    pre_parsed_items = None  # For formats we can parse directly

    if ext in ['txt', 'html', 'htm', 'eml']:
        # Text-based files - decode as text
        try:
            content_for_llm = file_data.decode('utf-8')
        except UnicodeDecodeError:
            content_for_llm = file_data.decode('latin-1')

    elif ext in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
        # Image files - send as base64 to vision model
        image_data = base64.standard_b64encode(file_data).decode('utf-8')
        media_type = f"image/{ext}" if ext != 'jpg' else 'image/jpeg'

    elif ext == 'ics':
        # ICS calendar files - parse directly
        try:
            pre_parsed_items = _parse_ics_file(file_data)
            if pre_parsed_items:
                return {
                    'success': True,
                    'items': pre_parsed_items,
                    'filename': filename
                }, 200
        except Exception as e:
            # Fall back to LLM parsing if direct parse fails
            try:
                content_for_llm = file_data.decode('utf-8')
            except UnicodeDecodeError:
                content_for_llm = file_data.decode('latin-1')

    elif ext == 'json':
        # JSON files - try to parse as trip data
        try:
            pre_parsed_items = _parse_json_trip(file_data)
            if pre_parsed_items:
                return {
                    'success': True,
                    'items': pre_parsed_items,
                    'filename': filename
                }, 200
        except Exception as e:
            # Fall back to LLM parsing
            try:
                content_for_llm = file_data.decode('utf-8')
            except UnicodeDecodeError:
                return {'error': 'Invalid JSON file encoding'}, 400

    elif ext in ['xlsx', 'xls']:
        # Excel files - extract as text table
        try:
            content_for_llm = _parse_excel_to_text(file_data, ext)
        except ImportError:
            return {'error': 'Excel processing not available. Please install openpyxl: pip install openpyxl'}, 500
        except Exception as e:
            return {'error': f'Error reading Excel file: {str(e)}'}, 400

    elif ext in ['docx', 'doc']:
        # Word documents - extract text
        try:
            content_for_llm = _parse_word_to_text(file_data, ext)
        except ImportError:
            return {'error': 'Word document processing not available. Please install python-docx: pip install python-docx'}, 500
        except Exception as e:
            return {'error': f'Error reading Word document: {str(e)}'}, 400

    elif ext == 'pdf':
        # PDF files - try to extract text, or use vision for scanned PDFs
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_data, filetype="pdf")
            text_content = []
            for page in doc:
                text_content.append(page.get_text())
            content_for_llm = "\n".join(text_content)

            # If no text extracted, it might be a scanned PDF - try first page as image
            if not content_for_llm.strip():
                page = doc[0]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                image_data = base64.standard_b64encode(pix.tobytes("png")).decode('utf-8')
                media_type = "image/png"

            doc.close()
        except ImportError:
            return {'error': 'PDF processing not available. Please install PyMuPDF: pip install pymupdf'}, 500
        except Exception as e:
            return {'error': f'Error reading PDF: {str(e)}'}, 400

    else:
        return {'error': f'Unsupported file type: {ext}'}, 400

    # Build the LLM prompt
    system_prompt = """You are a travel document parser. Extract travel-related items from the uploaded document.

For each item you find, extract:
- title: A clear name for the item (e.g., "LH 2416 MUC → ARN", "Hotel Duomo Firenze")
- category: One of: flight, transport, hotel, meal, activity, attraction, other
- date: The date if available (YYYY-MM-DD format)
- time: Start/departure time (HH:MM format, 24-hour)
- end_time: End/arrival time if available (HH:MM format, 24-hour) - IMPORTANT for flights and trains!
- location: City or address (destination for flights/trains)
- notes: Any additional relevant details (confirmation numbers, seat assignments, flight duration, etc.)

For FLIGHTS and TRAINS: Always extract both departure time (time) and arrival time (end_time) if shown.

Return your response as a JSON array of items. Example:
```json
[
  {
    "title": "LH 2416 MUC → ARN",
    "category": "flight",
    "date": "2025-12-17",
    "time": "12:10",
    "end_time": "14:25",
    "location": "Stockholm, Sweden",
    "notes": "Lufthansa, Airbus A321, Economy, 2h 15m nonstop"
  }
]
```

If you cannot extract any travel items, return an empty array: []
Only return the JSON array, no other text."""

    try:
        client = Anthropic()

        if image_data:
            # Use vision for images
            messages = [{
                'role': 'user',
                'content': [
                    {
                        'type': 'image',
                        'source': {
                            'type': 'base64',
                            'media_type': media_type,
                            'data': image_data
                        }
                    },
                    {
                        'type': 'text',
                        'text': f'Extract travel items from this document (filename: {filename})'
                    }
                ]
            }]
        else:
            # Use text content
            messages = [{
                'role': 'user',
                'content': f'Extract travel items from this document (filename: {filename}):\n\n{content_for_llm[:10000]}'  # Limit content
            }]

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2048,
            system=system_prompt,
            messages=messages
        )

        response_text = response.content[0].text.strip()

        # Parse the JSON response
        # Remove markdown code block if present
        if response_text.startswith('```'):
            lines = response_text.split('\n')
            response_text = '\n'.join(lines[1:-1] if lines[-1] == '```' else lines[1:])

        items = json.loads(response_text)

        if not isinstance(items, list):
            items = []

        return {
            'success': True,
            'items': items,
            'filename': filename
        }, 200

    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        print(f"Response was: {response_text}")
        return {'error': 'Failed to parse extracted items'}, 500
    except Exception as e:
        print(f"Upload plan error: {e}")
        import traceback
        traceback.print_exc()
        return {'error': f'Error processing file: {str(e)}'}, 500
